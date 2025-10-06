"""
Document processing pipeline for legal documents.
Handles PDF, DOCX, and TXT files with legal-specific preprocessing.
"""

import logging
import asyncio
import re
from typing import Dict, List, Optional, Any, Tuple, Union
from datetime import datetime
from pathlib import Path
import hashlib
import tempfile
import os

import pypdf
from docx import Document
import tiktoken

from app.core.config import settings

logger = logging.getLogger(__name__)


class DocumentProcessingError(Exception):
    """Raised when document processing fails."""
    pass


class LegalDocumentProcessor:
    """Processor for legal documents with specialized handling."""
    
    def __init__(self):
        self.supported_formats = settings.SUPPORTED_FORMATS.split(",")
        self.max_chunk_size = settings.MAX_CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
        
        # Legal-specific patterns
        self.citation_patterns = [
            r'\b\d+\s+[A-Z][a-z]+\.?\s+\d+\b',  # Case citations
            r'\b\d+\s+[A-Z]\.?\s*\d+d?\s+\d+\b',  # Court reports
            r'\b[A-Z]+\s+§\s*\d+(?:\.\d+)*\b',  # Statutes
            r'\bU\.S\.C\.\s*§\s*\d+\b',  # US Code
            r'\bC\.F\.R\.\s*§\s*\d+\b',  # Code of Federal Regulations
        ]
        
        self.legal_headers = [
            "WHEREAS", "THEREFORE", "NOW THEREFORE", "IN WITNESS WHEREOF",
            "WITNESSETH", "RECITALS", "DEFINITIONS", "TERMS AND CONDITIONS",
            "REPRESENTATIONS AND WARRANTIES", "INDEMNIFICATION",
            "GOVERNING LAW", "JURISDICTION", "ENTIRE AGREEMENT"
        ]
    
    async def process_file(
        self,
        file_path: Union[str, Path],
        document_type: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Process a document file and extract structured information."""
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise DocumentProcessingError(f"File not found: {file_path}")
            
            # Determine file type
            file_extension = file_path.suffix.lower().lstrip('.')
            if file_extension not in self.supported_formats:
                raise DocumentProcessingError(f"Unsupported file format: {file_extension}")
            
            # Extract text based on file type
            if file_extension == "pdf":
                text = await self._extract_pdf_text(file_path)
            elif file_extension == "docx":
                text = await self._extract_docx_text(file_path)
            elif file_extension == "txt":
                text = await self._extract_txt_text(file_path)
            else:
                raise DocumentProcessingError(f"No handler for file type: {file_extension}")
            
            # Process the extracted text
            processed_data = await self._process_text(
                text=text,
                source_file=str(file_path),
                document_type=document_type,
                metadata=metadata
            )
            
            # Add file-specific metadata
            processed_data["metadata"].update({
                "file_path": str(file_path),
                "file_size": file_path.stat().st_size,
                "file_extension": file_extension,
                "processed_at": datetime.utcnow().isoformat()
            })
            
            return processed_data
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            raise DocumentProcessingError(f"Failed to process file: {e}")
    
    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            loop = asyncio.get_event_loop()
            
            def extract_text():
                text_parts = []
                with open(file_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file)
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                # Add page marker for reference
                                text_parts.append(f"\n--- PAGE {page_num + 1} ---\n")
                                text_parts.append(page_text)
                        except Exception as e:
                            logger.warning(f"Error extracting page {page_num + 1}: {e}")
                            continue
                
                return "\n".join(text_parts)
            
            text = await loop.run_in_executor(None, extract_text)
            
            if not text.strip():
                raise DocumentProcessingError("No text extracted from PDF")
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise DocumentProcessingError(f"PDF extraction failed: {e}")
    
    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            loop = asyncio.get_event_loop()
            
            def extract_text():
                doc = Document(file_path)
                text_parts = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            text_parts.append(row_text)
                
                return "\n".join(text_parts)
            
            text = await loop.run_in_executor(None, extract_text)
            
            if not text.strip():
                raise DocumentProcessingError("No text extracted from DOCX")
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise DocumentProcessingError(f"DOCX extraction failed: {e}")
    
    async def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            raise DocumentProcessingError("Unable to decode text file with any encoding")
            
        except Exception as e:
            logger.error(f"Error extracting TXT text: {e}")
            raise DocumentProcessingError(f"TXT extraction failed: {e}")
    
    async def _process_text(
        self,
        text: str,
        source_file: str,
        document_type: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Process extracted text and generate structured data."""
        try:
            # Clean and normalize text
            cleaned_text = self._clean_text(text)
            
            # Detect document type if not provided
            if not document_type:
                document_type = self._detect_document_type(cleaned_text)
            
            # Extract metadata from text
            extracted_metadata = self._extract_text_metadata(cleaned_text)
            
            # Combine metadata
            combined_metadata = {
                "source": source_file,
                "document_type": document_type,
                "character_count": len(cleaned_text),
                "word_count": len(cleaned_text.split()),
                **extracted_metadata
            }
            
            if metadata:
                combined_metadata.update(metadata)
            
            # Generate document hash for deduplication
            doc_hash = hashlib.md5(cleaned_text.encode()).hexdigest()
            combined_metadata["document_hash"] = doc_hash
            
            return {
                "text": cleaned_text,
                "metadata": combined_metadata,
                "document_type": document_type,
                "source": source_file
            }
            
        except Exception as e:
            logger.error(f"Error processing text: {e}")
            raise DocumentProcessingError(f"Text processing failed: {e}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content."""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove page headers/footers (common patterns)
        text = re.sub(r'\n--- PAGE \d+ ---\n', '\n', text)
        text = re.sub(r'\nPage \d+( of \d+)?\n', '\n', text)
        
        # Fix common OCR errors in legal documents
        text = re.sub(r'\bSection\s+(\d+)', r'Section \1', text)
        text = re.sub(r'\bArticle\s+([IVX]+)', r'Article \1', text)
        
        # Preserve legal citations and references
        text = re.sub(r'(\d+)\s+([A-Z][a-z]+\.?)\s+(\d+)', r'\1 \2 \3', text)
        
        return text.strip()
    
    def _detect_document_type(self, text: str) -> str:
        """Detect the type of legal document."""
        text_lower = text.lower()
        
        # Contract indicators
        contract_indicators = [
            "agreement", "contract", "whereas", "party", "parties",
            "consideration", "terms and conditions", "hereby agree"
        ]
        
        # Brief indicators
        brief_indicators = [
            "brief", "motion", "memorandum", "court", "plaintiff",
            "defendant", "respectfully submitted", "comes now"
        ]
        
        # Will/Estate indicators
        estate_indicators = [
            "last will", "testament", "executor", "beneficiary",
            "estate", "bequest", "devise", "probate"
        ]
        
        # Corporate indicators
        corporate_indicators = [
            "articles of incorporation", "bylaws", "merger",
            "acquisition", "board of directors", "shareholder"
        ]
        
        # Patent indicators
        patent_indicators = [
            "patent", "invention", "claim", "specification",
            "prior art", "embodiment"
        ]
        
        # Count indicators
        scores = {
            "contract": sum(1 for indicator in contract_indicators if indicator in text_lower),
            "brief": sum(1 for indicator in brief_indicators if indicator in text_lower),
            "estate": sum(1 for indicator in estate_indicators if indicator in text_lower),
            "corporate": sum(1 for indicator in corporate_indicators if indicator in text_lower),
            "patent": sum(1 for indicator in patent_indicators if indicator in text_lower)
        }
        
        # Return type with highest score
        if max(scores.values()) > 0:
            return max(scores, key=scores.get)
        
        return "general"
    
    def _extract_text_metadata(self, text: str) -> Dict[str, Any]:
        """Extract metadata from document text."""
        metadata = {}
        
        # Extract dates
        date_patterns = [
            r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b',
            r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
            r'\b\d{4}-\d{2}-\d{2}\b'
        ]
        
        dates = []
        for pattern in date_patterns:
            dates.extend(re.findall(pattern, text, re.IGNORECASE))
        
        if dates:
            metadata["dates_found"] = dates[:5]  # Limit to first 5 dates
        
        # Extract legal citations
        citations = []
        for pattern in self.citation_patterns:
            citations.extend(re.findall(pattern, text))
        
        if citations:
            metadata["legal_citations"] = citations[:10]  # Limit to first 10 citations
        
        # Extract party names (basic pattern)
        party_patterns = [
            r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),?\s+(?:Plaintiff|Defendant|Petitioner|Respondent)\b',
            r'\b(?:Plaintiff|Defendant|Petitioner|Respondent):?\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\b'
        ]
        
        parties = []
        for pattern in party_patterns:
            parties.extend(re.findall(pattern, text))
        
        if parties:
            metadata["parties"] = list(set(parties))[:5]  # Unique parties, limit to 5
        
        # Extract jurisdiction indicators
        jurisdiction_patterns = [
            r'\b(United States District Court)\b',
            r'\b(Supreme Court)\b',
            r'\b(Court of Appeals)\b',
            r'\b([A-Z][a-z]+\s+(?:County|Parish))\b',
            r'\b([A-Z][a-z]+\s+State)\b'
        ]
        
        jurisdictions = []
        for pattern in jurisdiction_patterns:
            jurisdictions.extend(re.findall(pattern, text, re.IGNORECASE))
        
        if jurisdictions:
            metadata["jurisdictions"] = list(set(jurisdictions))[:3]
        
        # Legal document structure indicators
        structure_indicators = {
            "has_whereas_clauses": bool(re.search(r'\bWHEREAS\b', text, re.IGNORECASE)),
            "has_definitions": bool(re.search(r'\bDEFINITIONS?\b', text, re.IGNORECASE)),
            "has_signatures": bool(re.search(r'\b(?:SIGNATURE|EXECUTED|WITNESSED)\b', text, re.IGNORECASE)),
            "has_exhibits": bool(re.search(r'\bEXHIBIT\s+[A-Z0-9]\b', text, re.IGNORECASE))
        }
        
        metadata.update(structure_indicators)
        
        return metadata
    
    def get_token_count(self, text: str, model: str = "gpt-3.5-turbo") -> int:
        """Get token count for text using tiktoken."""
        try:
            encoding = tiktoken.encoding_for_model(model)
            return len(encoding.encode(text))
        except Exception:
            # Fallback to simple word count estimation
            return len(text.split()) * 1.3  # Rough approximation
    
    def estimate_processing_time(self, file_path: Union[str, Path]) -> float:
        """Estimate processing time based on file size."""
        file_path = Path(file_path)
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        
        # Rough estimates (seconds per MB)
        time_estimates = {
            "pdf": 2.0,
            "docx": 1.0,
            "txt": 0.1
        }
        
        file_extension = file_path.suffix.lower().lstrip('.')
        base_time = time_estimates.get(file_extension, 1.0)
        
        return file_size_mb * base_time
    
    async def batch_process_files(
        self,
        file_paths: List[Union[str, Path]],
        batch_size: int = None,
        progress_callback: Optional[callable] = None
    ) -> List[Dict[str, Any]]:
        """Process multiple files in batches."""
        batch_size = batch_size or settings.PROCESSING_BATCH_SIZE
        results = []
        total_files = len(file_paths)
        
        for i in range(0, total_files, batch_size):
            batch = file_paths[i:i + batch_size]
            batch_results = []
            
            # Process batch concurrently
            tasks = [self.process_file(file_path) for file_path in batch]
            
            try:
                batch_results = await asyncio.gather(*tasks, return_exceptions=True)
            except Exception as e:
                logger.error(f"Error in batch processing: {e}")
                batch_results = [DocumentProcessingError(str(e)) for _ in batch]
            
            # Handle results and exceptions
            for j, result in enumerate(batch_results):
                if isinstance(result, Exception):
                    logger.error(f"Failed to process {batch[j]}: {result}")
                    results.append({
                        "error": str(result),
                        "file_path": str(batch[j]),
                        "processed_at": datetime.utcnow().isoformat()
                    })
                else:
                    results.append(result)
            
            # Progress callback
            if progress_callback:
                progress = min(i + batch_size, total_files) / total_files
                await progress_callback(progress, f"Processed {len(results)}/{total_files} files")
        
        return results


# Global document processor instance
document_processor = LegalDocumentProcessor()